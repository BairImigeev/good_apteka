import os
from docx import Document
import re
import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

source_template = '\\\\nasdrkb2\\Сетевая2\\Обменник\\!МИС\\Bair\\template\\00000005900.xml'

# source_in = '\\nasdrkb2\\Сетевая2\\Обменник\\Аптека\\ДОГОВОРЫ 2023\\спецификации для хмл'

source_in = '\\nasdrkb2\\Сетевая2\\Обменник\\Аптека\\ДОГОВОРЫ 2023\\спецификации для хмл\контракты ЛП 2023\\77 БФК инсулин'

month = {'январ': '01', 'феврал': '02', 'март': '03', 'апрел': '04', 'мая': '05', 'июн': '06',
         'июл': '07', 'август': '08', 'сентябр': '09', 'октябр': '10', 'ноябр': '11', 'декабр': '12'}


def get_date(date_dogovor):
    for i in month.keys():
        print(i)
        if i in date_dogovor:
            print(i)
            date_val = date_dogovor.replace(i, str(month[i]))
            date_val = re.sub('\D+', '', date_val)
            day_object = datetime.datetime.strptime(date_val, '%d%m%Y').strftime('%d.%m.%Y')
            return str(day_object)


def get_locate_file():
    locate_files = []
    search_files =[]
    for rootdir, dirs, files in os.walk(source_in):
        for file in files:
            if (file.split('.')[-1]) =='docx':
                locate = os.path.join(rootdir, file)
                locate_files.append(rootdir)
                search_files.append(locate)
    return locate_files, search_files


def get_date_doc(doc):
    for j in doc.paragraphs:
        formatting = j.paragraph_format
        date_doc = re.search('«\d\d»\s\w+\s\d\d\d\d' , j.text)
        if date_doc and formatting.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            date_dogovor = re.sub("['«','»']", '', date_doc[0])
            # print(date_dogovor)
            for i in month.keys():
                if i in date_dogovor:
                    date_val = date_dogovor.replace(i, str(month[i]))
                    date_val = re.sub("[\D+',' ']",'', date_val)
                    date_val = datetime.datetime.strptime(date_val, '%d%m%Y').strftime('%d.%m.%Y')
                    print(date_val)
                    return date_val


def supply_period(doc):
    for j in doc.paragraphs:
        date_end_postavki = re.search('3.2. Поставщик обязуется передать', j.text)
        date_end_postavki_2 = re.search('3.2. Срок поставки товара: ', j.text)

        if date_end_postavki or date_end_postavki_2:
            date_end_postavki = re.search('в течени\w\s\d+\s[(]\w+[)]\s\w+\sдн\w+', j.text)
            date_end_postavki = re.search('\d+', date_end_postavki[0])[0]
            return date_end_postavki


def get_date_end_dat(doc):
    for j in doc.paragraphs:
        date_end_srok_postavki = re.search('12.1. Договор вступает в силу с даты его подписания и действует до ', j.text)
        if date_end_srok_postavki:
            date_doc_end_srok = re.search('«\d\d»\s\w+\s\d\d\d\d', j.text)
            date_end_dogovor = re.sub("['«','»']", '', date_doc_end_srok[0])
            for i in month.keys():
                if i in date_end_dogovor:
                    date_end = date_end_dogovor.replace(i, str(month[i]))
                    date_end = re.sub("[\D+',' ']",'', date_end)
                    date_end = datetime.datetime.strptime(date_end, '%d%m%Y').strftime('%d.%m.%Y')
                    return date_end


def get_company(doc, source):
    for table in doc.tables:
        for index, row in enumerate(table.rows):
            if index == 0:
                row_text = list(cell.text for cell in row.cells)
                if re.search('«Детская', row_text[0]):
                    words = re.findall(r'(«\w+\W+\w+»)', row_text[1]) or re.findall(r'(«\w+\w+\w+»)', row_text[1])
                    if words:
                        # print(source)
                        # print('компания v1: ', words[0][1:-1])
                        return words[0][1:-1]
                    words = re.search('«\w+\s+\w+', row_text[1])
                    words_2 = re.search('\w+»', row_text[1])
                    if words and words_2:
                        words = re.sub("['«','»']", '', words[0])
                        words_2 = re.sub("['«','»']", '', words_2[0])
                        words = words + ' ' + words_2
                        return words


def main():
    # located = get_locate_file()[0]
    located_file = get_locate_file()[1]
    print(located_file)
    for i in located_file:
        # print(i)
        doc = Document(i)
        print(i)
        print(get_date_end_dat(doc))
        print(supply_period(doc))
        print(get_company(doc, i))


if __name__ == '__main__':
    main()