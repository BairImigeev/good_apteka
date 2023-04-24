import encodings.utf_8
import os
from docx import Document
import xml.etree.ElementTree as ET
import datetime
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

source_template = '\\\\nasdrkb2\\Сетевая2\\Обменник\\!МИС\\Bair\\template\\00000005900.xml'
source_in = '\\\\nasdrkb2\\Сетевая2\\Обменник\\Аптека\\ДОГОВОРЫ 2023\\спецификации для хмл'
# source_in = 'C:\\apteka\\'

month = {'январ': '01', 'феврал': '02', 'март': '03', 'апрел': '04', 'мая':'05', 'июн': '06', 'июл': '07',
          'август': '08', 'сентябр': '09', 'октябр': '10', 'ноябр': '11', 'декабр': '12'}

a = 'ns0'
b = 'ns1'
c = 'ns2'
d = 'dt="uuid:BDC6E3F0-6DA3-11d1-A2A3-00AA00C14882"'
e = ' /'


def get_xml():
    tree = ET.parse(source_template)  # Parse XML
    return tree


def get_locate_file():
    locate_files = []
    search_files =[]
    for rootdir, dirs, files in os.walk(source_in):
        for file in files:
            if (file.split('.')[-1]) =='docx':
                locate = os.path.join(rootdir, file)
                locate_files.append(rootdir)
                search_files.append(locate)
    return locate_files, search_files


def give_element_xml(tree):
    root = tree.getroot()
    new_element = root[1][0]
    return new_element


def get_table(table, priznak_table):
    list_data = []
    for index, row in enumerate(table.rows[1:-1]):

        data = {}
        data['COMPANY_FROM_ID'] = ''
        data['DEP_TO_ID'] = ''
        data['MAN_CONFIRM_ID'] = ''
        data['TYPE'] = ''
        data['CODE'] = ''
        data['CODE2'] = ''
        data['DAT'] = ''
        data['DAT_CONFIRM'] = ''
        if priznak_table == 0:
            data['PROD_NOTE'] = row.cells[2].text
            data['QTY'] = row.cells[5].text.replace(' ', '')
            data['QTYUNIT'] = row.cells[5].text.replace(' ', '')
            data['PRICE'] = str(
                float(row.cells[7].text.replace('\xa0', '').replace(' ', '').replace(',', '.'))).replace('.', ',')
            data['PRICE_NDS'] = str(
                float(row.cells[7].text.replace('\xa0', '').replace(' ', '').replace(',', '.'))).replace('.', ',')
            data['PRODUCT_NAME'] = row.cells[1].text
            data['FIRM'] = get_firm(row.cells[3].text)
            data['FIRM_CITY'] = get_firm_city(row.cells[3].text)
            data['UNIT'] = row.cells[4].text
        else:
            data['PROD_NOTE'] = row.cells[3].text
            data['QTY'] = row.cells[6].text.replace(' ', '')
            data['QTYUNIT'] = row.cells[6].text.replace(' ', '')
            data['PRICE'] = str(
                float(row.cells[8].text.replace('\xa0', '').replace(' ', '').replace(',', '.'))).replace('.', ',')
            data['PRICE_NDS'] = str(
                float(row.cells[8].text.replace('\xa0', '').replace(' ', '').replace(',', '.'))).replace('.', ',')
            data['PRODUCT_NAME'] = row.cells[2].text
            data['FIRM'] = get_firm(row.cells[4].text)
            data['FIRM_CITY'] = get_firm_city(row.cells[4].text)
            data['UNIT'] = row.cells[5].text

        data['NOTE'] = ''
        data['PAYTYPE_ID'] = ''
        data['NDS'] = ''
        data['DAT2'] = ''
        data['END_DAT'] = ''
        data['BUY_TYPE_ID'] = ''
        data['SUPPLY_PERIOD'] = ''
        data['PRODUCT_TYPE_ID'] = ''
        data['MOL_TO_ID'] = ''
        data['NUM_DOC'] = ''
        data['PRODUCT_CODE'] = ''
        data['LIFE_DAT'] = ''
        data['SER'] = ''
        data['PRODUCT_NAME_ID'] = ''
        data['MNN'] = ''
        data['MNN_R'] = ''
        data['PRODUCT_NAME_LAT'] = ''
        data['UNIT_PAT'] = ''
        data['PRODUCT_FORM'] = ''
        data['PRODUCT_DOSAGE'] = ''
        data['PRODUCT_FULLNAME'] = ''

        list_data.append(data)

    return list_data


def get_firm(cell_from_table):
    # print(cell_from_table)
    firm = cell_from_table.split(',')
    firm = firm[0]
    # print(firm)
    return firm


def get_firm_city(cell_from_table):
    firm_city = cell_from_table.split(',')
    firm_city = firm_city[1]
    # print(firm_city)
    return firm_city


def get_date_doc(doc):
    for j in doc.paragraphs:
        formatting = j.paragraph_format
        date_doc = re.search('«\d\d»\s\w+\s\d\d\d\d' , j.text)

        if date_doc and formatting.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            # print(date_doc)
            date_dogovor = re.sub("['«','»']", '', date_doc[0])
            # print(date_dogovor)
            for i in month.keys():
                if i in date_dogovor:
                    date_val = date_dogovor.replace(i, str(month[i]))
                    date_val = re.sub("[\D+',' ']",'', date_val)
                    date_val = datetime.datetime.strptime(date_val, '%d%m%Y').strftime('%d.%m.%Y')
                    # print(date_val)
                    return date_val


def get_date_end_dat(doc):
    for j in doc.paragraphs:
        date_end_srok_postavki = re.search('12.1. Договор вступает в силу с даты его подписания и действует до ', j.text)
        if date_end_srok_postavki:
            date_doc_end_srok = re.search('«\d\d»\s\w+\s\d\d\d\d', j.text)
            # print(date_doc_end_srok)
            date_end_dogovor = re.sub("['«','»']", '', date_doc_end_srok[0])
            for i in month.keys():
                if i in date_end_dogovor:
                    date_end = date_end_dogovor.replace(i, str(month[i]))
                    date_end = re.sub("[\D+',' ']",'', date_end)
                    date_end = datetime.datetime.strptime(date_end, '%d%m%Y').strftime('%d.%m.%Y')
                    return date_end


def supply_period(doc):
    for j in doc.paragraphs:
        date_end_postavki = re.search('3.2. Поставщик обязуется передать', j.text)
        date_end_postavki_2 = re.search('3.2. Срок поставки товара: ', j.text)

        if date_end_postavki or date_end_postavki_2:
            date_end_postavki = re.search('в течени\w\s\d+\s[(]\w+[)]\s\w+\sдн\w+', j.text)
            date_end_postavki_3 = re.search('в течени\w\s\d+\s\w+\sдн\w+', j.text)
            if date_end_postavki:
                date_end_postavki = re.search('\d+', date_end_postavki[0])[0]
                return date_end_postavki
            if date_end_postavki_3:
                date_end_postavki = re.search('\d+', date_end_postavki_3[0])[0]
                return date_end_postavki


def get_company(doc):
    for table in doc.tables:
        for index, row in enumerate(table.rows):
            if index == 0:
                row_text = list(cell.text for cell in row.cells)
                if re.search('«Детская', row_text[0]):
                    words_v1 = re.findall(r'(«\w+\W+\w+»)', row_text[1]) or re.findall(r'(«\w+\w+\w+»)', row_text[1])
                    # words_v2 = re.findall(r'')
                    if words_v1:
                        # print(source)
                        # print('компания v1: ', words_v1[0][1:-1])
                        return words_v1[0][1:-1]
                    words_v1 = re.search('«\w+\s+\w+', row_text[1])
                    words_2 = re.search('\w+»', row_text[1])
                    if words_v1 and words_2:
                        words = re.sub("['«','»']", '', words_v1[0])
                        words_2 = re.sub("['«','»']", '', words_2[0])
                        words = words + ' ' + words_2
                        return words


def get_num(doc):
    for j in doc.paragraphs:
        if re.search('ДОГОВОР № ', j.text):
            numberdoc = j.text.replace('ДОГОВОР № ', '')
            return numberdoc
        if re.search('Договор № ', j.text):
            numberdoc = j.text.replace('Договор № ', '')
            return numberdoc
        if re.search('Договору № \d{4}.\d+', j.text):
            numberdoc = j.text.replace('к Договору № ', '')
            return numberdoc


def get_data(doc):
    for table in doc.tables:
        for index, row in enumerate(table.rows):
            if index == 0:
                # print('ok')

                row_text = list(cell.text for cell in row.cells)
                if re.search('БИК', row_text[0]):
                    break
                if (re.search('МНН', row_text[1]) or re.search('Наименование товара', row_text[1])
                        or re.search('Наименование', row_text[1])) \
                        and (re.search('Наименование товара', row_text[2])
                        or re.search('Торговое наименование', row_text[2])
                        or re.search('Наименование', row_text[2])):
                    # print('ok_1')
                    priznak_table = 1
                    if re.search('Характеристики товара', row_text[3]):
                        # print('ok_3')
                        if re.search('Страна', row_text[4]):
                            # print('ok_4')
                            if re.search('Ед. изм', row_text[5]) or re.search('Единица измерения', row_text[5]):
                                # print('ok_5')
                                if re.search('Кол-во', row_text[6]) or re.search('Количество', row_text[6]):
                                    # print('ok_6')
                                    if re.search('Цена за единицу', row_text[7]):
                                        print('она')
                                        data_prod = get_table(table, priznak_table)
                                        # print(data_prod)
                                        return data_prod

                if re.search('Наименование', row_text[1]):
                    priznak_table = 0
                    # print('ok_1_2')
                    # print(row_text[1])
                    if re.search('Характеристики товара', row_text[2]):
                        # print('ok_3')
                        if re.search('Страна', row_text[3]):
                            # print('ok_4')
                            if re.search('Ед. изм', row_text[4]) or re.search('Единица измерения', row_text[5]):
                                # print('ok_5')
                                if re.search('Кол-во', row_text[5]) or re.search('Количество', row_text[6]):
                                    # print('ok_6')
                                    if re.search('Цена за единицу', row_text[6]):
                                        print('она')
                                        data_prod = get_table(table, priznak_table)
                                        # print(data_prod)
                                        return data_prod


def main():
    located = get_locate_file()[0]
    located_file = get_locate_file()[1]

    for i in located_file:
        print(i)
        doc = Document(i)
        company = get_company(doc)
        print(' Компания : ', company)
        data_doc = get_date_doc(doc)
        print(' Дата документа : ', data_doc)
        number_doc = get_num(doc)
        print(' Номера документа : ', number_doc)
        supply_per = supply_period(doc)
        print(' Период : ', supply_per)
        end_dat = get_date_end_dat(doc)
        print(' Конечный срок поставки : ', end_dat)

        data_prod = get_data(doc)
        my_tree = get_xml()
        root = my_tree.getroot()
        copy_new_element = give_element_xml(my_tree)

        if data_prod is None:
            print('Файл : ')
            print(i)
            print('Ошибка. В данных None : ')
            print(copy_new_element.attrib)

        for j in data_prod:
            new = ET.Element('')
            new.tail = copy_new_element.tail
            new.tag = copy_new_element.tag

            j['END_DAT'] = end_dat
            j['SUPPLY_PERIOD'] = supply_per
            j['COMPANY_FROM_ID'] = company
            j['DAT2'] = data_doc
            j['NUM_DOC'] = number_doc
            j['PRODUCT_TYPE_ID'] = 'Медикаменты'
            if 'контракт' in str(i):
                print('контракт')
                j['BUY_TYPE_ID'] = '9802'
            else:
                j['BUY_TYPE_ID'] = '9801'

            new.attrib = j

            for key, value in new.attrib.items():
                if value is None:
                    print('Ошибка записи в файл, требуется посмотреть данные полей')
                    print(i)
                    print(key)
                    break
                else:
                    continue

            root[1].append(new)

        root[1].remove(root[1][0])
        my_tree.write(f'{i}.xml', encoding='utf-8')


if __name__ == '__main__':
    main()












